"""
将 AI 生成的建议追加到 Word 文档

用法:
    python scripts/append_ai_recommendations.py <docx_path> <recommendations_file>

其中 recommendations_file 是一个包含 AI 建议的 markdown 文件。
"""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def append_recommendations(docx_path: str, recommendations: str) -> None:
    """
    将 AI 建议追加到 Word 文档

    Args:
        docx_path: Word 文档路径
        recommendations: AI 生成的建议文本 (markdown 格式)
    """
    doc = Document(docx_path)

    # 添加分隔
    doc.add_paragraph()

    # 添加 AI 分析标题
    heading = doc.add_heading("AI 智能分析与建议", level=1)

    # 添加说明
    note = doc.add_paragraph()
    note_run = note.add_run("以下内容由 Claude AI 基于交易数据自动生成：")
    note_run.italic = True
    note_run.font.size = Pt(10)

    doc.add_paragraph()

    # 解析并添加建议内容
    lines = recommendations.strip().split("\n")
    current_heading_level = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 处理 markdown 标题
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=2)
        elif line.startswith("- ") or line.startswith("* "):
            # 列表项
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. "):
            # 有序列表
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith("**") and line.endswith("**"):
            # 粗体段落
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        else:
            # 普通段落
            doc.add_paragraph(line)

    # 保存
    doc.save(docx_path)
    print(f"AI 建议已追加到: {docx_path}")


def main():
    if len(sys.argv) < 3:
        print("用法: python scripts/append_ai_recommendations.py <docx_path> <recommendations_file>")
        print("或者通过管道传入建议内容:")
        print("  echo '建议内容' | python scripts/append_ai_recommendations.py <docx_path> -")
        sys.exit(1)

    docx_path = sys.argv[1]
    recommendations_source = sys.argv[2]

    if recommendations_source == "-":
        # 从标准输入读取
        recommendations = sys.stdin.read()
    else:
        # 从文件读取
        recommendations = Path(recommendations_source).read_text(encoding="utf-8")

    append_recommendations(docx_path, recommendations)


if __name__ == "__main__":
    main()
