#!/usr/bin/env python3
"""
将投资教练建议追加到 Word 报告

用法:
    python scripts/append_coach_to_docx.py output/2025年美港股交易分析报告.docx /tmp/coach.md

或者交互式使用（从标准输入读取 markdown）:
    cat /tmp/coach.md | python scripts/append_coach_to_docx.py output/2025年美港股交易分析报告.docx
"""

import sys
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def parse_markdown(content: str) -> list[dict]:
    """
    解析 Markdown 内容为结构化数据

    Returns:
        list of {type: heading/paragraph/list_item, level: int, content: str}
    """
    elements = []
    lines = content.split('\n')

    for line in lines:
        line = line.rstrip()

        if not line:
            continue

        # 标题
        if line.startswith('#'):
            match = re.match(r'^(#+)\s+(.+)$', line)
            if match:
                level = len(match.group(1))
                elements.append({
                    'type': 'heading',
                    'level': level,
                    'content': match.group(2)
                })
                continue

        # 列表项
        if line.startswith('- ') or line.startswith('* ') or line.startswith('□ '):
            content = line[2:].strip()
            elements.append({
                'type': 'list_item',
                'content': content,
                'checkbox': line.startswith('□')
            })
            continue

        # 编号列表
        if re.match(r'^\d+\.\s+', line):
            content = re.sub(r'^\d+\.\s+', '', line)
            elements.append({
                'type': 'numbered_item',
                'content': content
            })
            continue

        # 普通段落
        elements.append({
            'type': 'paragraph',
            'content': line
        })

    return elements


def append_to_docx(docx_path: Path, markdown_content: str) -> None:
    """
    将 Markdown 内容追加到 Word 文档

    Args:
        docx_path: Word 文档路径
        markdown_content: Markdown 格式的内容
    """
    doc = Document(docx_path)

    # 添加分页符
    doc.add_page_break()

    # 解析 Markdown
    elements = parse_markdown(markdown_content)

    for elem in elements:
        if elem['type'] == 'heading':
            level = elem['level']
            content = elem['content']

            # 去除 emoji 用于 heading（保留在文本中）
            heading = doc.add_heading(content, level=min(level, 3))

            # 设置中文字体
            for run in heading.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

                # 根据内容设置颜色
                if '⚠️' in content or '核心问题' in content:
                    run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)  # 红色
                elif '✅' in content or '做得好' in content:
                    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # 绿色
                elif '💡' in content or '建议' in content:
                    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)  # 蓝色

        elif elem['type'] == 'list_item':
            para = doc.add_paragraph()
            if elem.get('checkbox'):
                para.add_run('□ ')
            else:
                para.add_run('• ')
            run = para.add_run(elem['content'])
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        elif elem['type'] == 'numbered_item':
            para = doc.add_paragraph(elem['content'], style='List Number')
            for run in para.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        elif elem['type'] == 'paragraph':
            content = elem['content']

            # 跳过代码块标记
            if content.startswith('```'):
                continue

            para = doc.add_paragraph()
            run = para.add_run(content)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(11)

            # 特殊样式
            if content.startswith('**') and content.endswith('**'):
                run.font.bold = True
            if '框架参考' in content or 'V10.10' in content:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 保存
    doc.save(docx_path)
    print(f"✓ 投资教练建议已追加到: {docx_path}")


def main():
    if len(sys.argv) < 2:
        print("用法: python scripts/append_coach_to_docx.py <docx_path> [markdown_file]")
        print("  或: cat markdown.md | python scripts/append_coach_to_docx.py <docx_path>")
        sys.exit(1)

    docx_path = Path(sys.argv[1])

    if not docx_path.exists():
        print(f"错误: 文件不存在 - {docx_path}")
        sys.exit(1)

    # 读取 Markdown 内容
    if len(sys.argv) >= 3:
        md_path = Path(sys.argv[2])
        if not md_path.exists():
            print(f"错误: Markdown 文件不存在 - {md_path}")
            sys.exit(1)
        markdown_content = md_path.read_text(encoding='utf-8')
    else:
        # 从标准输入读取
        markdown_content = sys.stdin.read()

    if not markdown_content.strip():
        print("错误: 没有内容可追加")
        sys.exit(1)

    append_to_docx(docx_path, markdown_content)


if __name__ == "__main__":
    main()
