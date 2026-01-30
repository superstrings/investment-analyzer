#!/usr/bin/env python3
"""
更新 Word 报告的"结论与建议"章节

使用 pandoc 将 Markdown 转换为 Word，然后嵌入到主报告中。
这种方法比手动解析 Markdown 排版效果更好。

用法:
    python scripts/update_docx_conclusion.py output/2025年美港股交易分析报告.docx output/2025年投资教练点评.md

或者从标准输入读取:
    cat output/coach.md | python scripts/update_docx_conclusion.py output/2025年美港股交易分析报告.docx
"""

import subprocess
import sys
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


def find_conclusion_section(doc: Document) -> int:
    """
    找到"结论与建议"章节的起始位置

    Returns:
        段落索引，如果找不到返回 -1
    """
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # 匹配新格式 "第三部分：结论与建议" 或 "第三部分: 结论与建议"
        if re.match(r'^第三部分[：:]\s*结论与建议', text):
            return i
        # 匹配旧格式 "十、结论与建议" 或 "十一、结论与建议" 等
        if re.match(r'^(十[一二三四五六七八九十]*)?[、．.]?\s*结论与建议', text):
            return i
        # 也匹配英文格式
        if 'Conclusion' in text or '结论' in text:
            if para.style and 'Heading' in para.style.name:
                return i
    return -1


def find_next_heading(doc: Document, start_idx: int) -> int:
    """
    找到下一个标题的位置

    Returns:
        段落索引，如果没有更多标题返回文档末尾
    """
    for i in range(start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.style and 'Heading' in para.style.name:
            # 检查是否是一级标题
            if para.style.name == 'Heading 1':
                return i
            # 或者匹配中文章节格式
            if re.match(r'^[一二三四五六七八九十]+[、．.]', para.text.strip()):
                return i
    return len(doc.paragraphs)


def convert_md_to_docx(markdown_content: str, output_path: Path) -> bool:
    """
    使用 pandoc 将 Markdown 转换为 Word 文档

    Args:
        markdown_content: Markdown 内容
        output_path: 输出 docx 文件路径

    Returns:
        是否成功
    """
    try:
        # 创建临时 md 文件
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write(markdown_content)
            temp_md = f.name

        # 使用 pandoc 转换
        result = subprocess.run(
            ['pandoc', temp_md, '-o', str(output_path), '--from=markdown', '--to=docx'],
            capture_output=True,
            text=True
        )

        # 清理临时文件
        Path(temp_md).unlink()

        if result.returncode != 0:
            print(f"pandoc 错误: {result.stderr}")
            return False

        return True
    except FileNotFoundError:
        print("错误: pandoc 未安装，请运行 brew install pandoc")
        return False
    except Exception as e:
        print(f"转换错误: {e}")
        return False


def merge_documents(main_doc_path: Path, coach_doc_path: Path, start_idx: int, end_idx: int) -> None:
    """
    将教练文档内容合并到主报告的指定位置

    Args:
        main_doc_path: 主报告路径
        coach_doc_path: 教练点评文档路径
        start_idx: 章节标题索引
        end_idx: 下一章节索引
    """
    main_doc = Document(main_doc_path)
    coach_doc = Document(coach_doc_path)

    # 删除旧内容（保留章节标题）
    paragraphs_to_remove = []
    for i in range(start_idx + 1, end_idx):
        if i < len(main_doc.paragraphs):
            paragraphs_to_remove.append(main_doc.paragraphs[i]._p)

    for p in paragraphs_to_remove:
        p.getparent().remove(p)

    # 保存并重新加载
    main_doc.save(main_doc_path)
    main_doc = Document(main_doc_path)

    # 重新找到标题位置
    start_idx = find_conclusion_section(main_doc)
    if start_idx == -1:
        print("错误: 重新加载后未找到章节")
        return

    title_para = main_doc.paragraphs[start_idx]

    # 添加框架版本说明
    framework_note = main_doc.add_paragraph()
    run = framework_note.add_run("（基于投资分析框架 V10.10，由 AI 投资教练生成）")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True

    # 移动到标题后
    framework_note._p.getparent().remove(framework_note._p)
    title_para._p.addnext(framework_note._p)

    current_element = framework_note._p

    # 复制教练文档的所有内容
    skip_first_heading = True  # 跳过第一个标题（投资教练点评）

    for element in coach_doc.element.body:
        # 跳过文档属性等非内容元素
        if element.tag.endswith('sectPr'):
            continue

        # 检查是否是第一个标题段落
        if skip_first_heading and element.tag.endswith('p'):
            para_text = ''.join(t.text or '' for t in element.iter() if t.text)
            if '投资教练点评' in para_text or para_text.strip().startswith('#'):
                skip_first_heading = False
                continue

        # 复制元素
        import copy
        new_element = copy.deepcopy(element)
        current_element.addnext(new_element)
        current_element = new_element

    # 保存
    main_doc.save(main_doc_path)


def update_conclusion(docx_path: Path, markdown_content: str) -> None:
    """
    更新 Word 文档的"结论与建议"章节

    Args:
        docx_path: Word 文档路径
        markdown_content: Markdown 格式的建议内容
    """
    doc = Document(docx_path)

    # 找到"结论与建议"章节
    start_idx = find_conclusion_section(doc)
    if start_idx == -1:
        print("警告: 未找到'结论与建议'章节，将追加到文档末尾")
        doc.add_page_break()
        doc.add_heading("结论与建议", level=1)
        doc.save(docx_path)
        doc = Document(docx_path)
        start_idx = find_conclusion_section(doc)

    # 找到下一个章节
    end_idx = find_next_heading(doc, start_idx)

    # 使用 pandoc 转换 markdown 为 docx
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx = Path(f.name)

    if not convert_md_to_docx(markdown_content, temp_docx):
        print("警告: pandoc 转换失败，使用简单文本模式")
        # 回退到简单模式
        _simple_update(docx_path, markdown_content, start_idx, end_idx)
        return

    # 合并文档
    merge_documents(docx_path, temp_docx, start_idx, end_idx)

    # 清理临时文件
    temp_docx.unlink()

    print(f"✓ 已更新报告的'结论与建议'章节: {docx_path}")


def _simple_update(docx_path: Path, markdown_content: str, start_idx: int, end_idx: int) -> None:
    """简单文本更新模式（pandoc 不可用时的回退方案）"""
    doc = Document(docx_path)

    # 删除旧内容
    paragraphs_to_remove = []
    for i in range(start_idx + 1, end_idx):
        if i < len(doc.paragraphs):
            paragraphs_to_remove.append(doc.paragraphs[i]._p)

    for p in paragraphs_to_remove:
        p.getparent().remove(p)

    doc.save(docx_path)
    doc = Document(docx_path)

    start_idx = find_conclusion_section(doc)
    title_para = doc.paragraphs[start_idx]

    # 添加简单文本
    lines = markdown_content.split('\n')
    current = title_para

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('#'):
            # 标题
            text = re.sub(r'^#+\s*', '', line)
            new_para = doc.add_paragraph()
            run = new_para.add_run(text)
            run.font.bold = True
            run.font.size = Pt(14)
        else:
            # 普通段落
            new_para = doc.add_paragraph(line)

        new_para._p.getparent().remove(new_para._p)
        current._p.addnext(new_para._p)
        current = new_para

    doc.save(docx_path)
    print(f"✓ 已更新报告的'结论与建议'章节 (简单模式): {docx_path}")


def main():
    if len(sys.argv) < 2:
        print("用法: python scripts/update_docx_conclusion.py <docx_path> [markdown_file]")
        print("  或: cat markdown.md | python scripts/update_docx_conclusion.py <docx_path>")
        sys.exit(1)

    docx_path = Path(sys.argv[1])

    if not docx_path.exists():
        print(f"错误: 文件不存在 - {docx_path}")
        sys.exit(1)

    # 读取 Markdown 内容
    if len(sys.argv) >= 3:
        md_path = Path(sys.argv[2])
        if not md_path.exists():
            print(f"错误: Markdown 文件不存在 - {md_path}")
            sys.exit(1)
        markdown_content = md_path.read_text(encoding='utf-8')
    else:
        # 从标准输入读取
        markdown_content = sys.stdin.read()

    if not markdown_content.strip():
        print("错误: 没有内容")
        sys.exit(1)

    update_conclusion(docx_path, markdown_content)


if __name__ == "__main__":
    main()
