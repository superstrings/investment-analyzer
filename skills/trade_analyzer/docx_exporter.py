"""
Word Exporter - Word 报告导出模块

使用 python-docx 生成交易分析报告：
- 标题/章节结构
- 统计表格
- 嵌入图表（PNG 格式）
- 结论与改进建议
"""

import io
from datetime import datetime
from decimal import Decimal
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .chart_generator import ChartGenerator
from .recommendation import InvestmentCoach, TradeRecommendation
from .statistics import TradeStatistics, StatisticsCalculator
from .trade_matcher import MatchedTrade


class DocxExporter:
    """Word 报告导出器"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self) -> None:
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles["Normal"]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(11)

    def export(
        self,
        trades: list[MatchedTrade],
        stats: TradeStatistics,
        charts: dict[str, bytes],
        output_path: Path,
        year: Optional[int] = None,
    ) -> Path:
        """
        导出交易分析报告到 Word

        Args:
            trades: 配对后的交易列表
            stats: 统计数据
            charts: 图表字典 {名称: PNG字节数据}
            output_path: 输出文件路径
            year: 年份

        Returns:
            输出文件路径
        """
        self.doc = Document()
        self._setup_styles()

        year = year or datetime.now().year

        # 标题
        title = self.doc.add_heading(f"{year}年美港股交易统计分析报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 报告日期
        date_para = self.doc.add_paragraph(
            f"报告生成日期：{datetime.now().strftime('%Y年%m月%d日')}"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 分离股票和期权交易
        stock_trades = [t for t in trades if not t.is_option]
        option_trades = [t for t in trades if t.is_option]

        # 一、整体交易表现
        self._add_overall_section(stats)

        # 二、盈亏统计
        self._add_profit_loss_section(stats)

        # 三、持仓时间分析
        self._add_holding_section(stats, charts)

        # 四、市场分布分析
        self._add_market_section(stats, charts)

        # 五、最佳交易 Top 5
        self._add_top_winners_section(stats)

        # 六、最大亏损 Top 5
        self._add_top_losers_section(stats)

        # 七、交易标的统计 Top 10
        self._add_stock_stats_section(stats)

        # 八、盈亏率分布
        self._add_profit_loss_distribution_section(stats, charts)

        # 九、月度盈亏趋势
        self._add_monthly_section(stats, charts)

        # 十、期权交易统计（如果有）
        if option_trades:
            self._add_option_section(stats, option_trades)

        # 十一、结论与建议（基于投资框架 V10.10 的智能建议）
        self._add_conclusion_section(stats, stock_trades, option_trades)

        # 保存文件
        self.doc.save(output_path)
        return output_path

    def _add_overall_section(self, stats: TradeStatistics) -> None:
        """添加整体交易表现章节"""
        self.doc.add_heading("一、整体交易表现", level=1)

        # 创建表格
        table = self.doc.add_table(rows=6, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["指标", "数值", "说明"]
        data = [
            ("总交易笔数", f"{stats.total_trades}笔", "已配对完成的买卖交易"),
            ("盈利笔数", f"{stats.winning_trades}笔", f"占比{stats.win_rate:.1%}"),
            ("亏损笔数", f"{stats.losing_trades}笔", f"占比{1-stats.win_rate:.1%}"),
            ("胜率", f"{stats.win_rate:.1%}", "盈利交易占总交易的比例"),
            ("盈亏比", f"{float(stats.profit_loss_ratio):.2f}", "平均盈利/平均亏损"),
        ]

        self._fill_table(table, headers, data)

    def _add_profit_loss_section(self, stats: TradeStatistics) -> None:
        """添加盈亏统计章节"""
        self.doc.add_heading("二、盈亏统计", level=1)

        table = self.doc.add_table(rows=6, cols=3)
        table.style = "Table Grid"

        headers = ["项目", "金额 (HKD)", "备注"]
        data = [
            ("总盈利", f"+{float(stats.total_profit):,.2f}", "所有盈利交易累计"),
            ("总亏损", f"-{float(stats.total_loss):,.2f}", "所有亏损交易累计"),
            ("净利润", f"{float(stats.net_profit):,.2f}", "总盈利-总亏损"),
            ("平均盈利", f"+{float(stats.avg_profit):,.2f}", "单笔盈利交易平均"),
            ("平均亏损", f"-{float(stats.avg_loss):,.2f}", "单笔亏损交易平均"),
        ]

        self._fill_table(table, headers, data)

        # 关键发现
        if stats.profit_loss_ratio < 1:
            finding = (
                f"关键发现：虽然胜率达到{stats.win_rate:.1%}，"
                f"但盈亏比仅为{float(stats.profit_loss_ratio):.2f}（低于1），"
                "说明平均亏损额大于平均盈利额。建议加强止损纪律，控制单笔亏损幅度。"
            )
            para = self.doc.add_paragraph(finding)
            para.runs[0].font.color.rgb = RGBColor(0x9C, 0x00, 0x06)

    def _add_holding_section(
        self, stats: TradeStatistics, charts: dict[str, bytes]
    ) -> None:
        """添加持仓时间分析章节"""
        self.doc.add_heading("三、持仓时间分析", level=1)

        table = self.doc.add_table(rows=5, cols=3)
        table.style = "Table Grid"

        headers = ["指标", "天数", "说明"]
        data = [
            ("平均持仓天数", f"{stats.avg_holding_days:.1f}天", "所有交易平均"),
            (
                "盈利交易平均持仓",
                f"{stats.avg_winning_holding_days:.1f}天",
                "持有时间较长",
            ),
            (
                "亏损交易平均持仓",
                f"{stats.avg_losing_holding_days:.1f}天",
                "持有时间较短",
            ),
            ("最长持仓", f"{stats.max_holding_days}天", ""),
        ]

        self._fill_table(table, headers, data)

        # 添加图表
        if "holding_days_hist" in charts and charts["holding_days_hist"]:
            self.doc.add_paragraph()
            self._add_image(charts["holding_days_hist"], "holding_days_hist")

    def _add_market_section(
        self, stats: TradeStatistics, charts: dict[str, bytes]
    ) -> None:
        """添加市场分布分析章节"""
        self.doc.add_heading("四、市场分布分析", level=1)

        if stats.market_stats:
            table = self.doc.add_table(rows=len(stats.market_stats) + 1, cols=5)
            table.style = "Table Grid"

            headers = ["市场", "交易笔数", "胜率", "净盈亏", "平均盈亏"]
            data = []
            for market, ms in stats.market_stats.items():
                market_name = {
                    "HK": "港股",
                    "US": "美股",
                    "SH": "沪市",
                    "SZ": "深市",
                }.get(market, market)
                data.append(
                    (
                        market_name,
                        f"{ms.total_trades}笔",
                        f"{ms.win_rate:.1%}",
                        f"{float(ms.net_profit):,.2f}",
                        f"{float(ms.avg_profit_loss):,.2f}",
                    )
                )

            self._fill_table(table, headers, data)

            # 分析文字
            if len(stats.market_stats) >= 2:
                markets_sorted = sorted(
                    stats.market_stats.items(),
                    key=lambda x: x[1].net_profit,
                    reverse=True,
                )
                top_market = markets_sorted[0]
                market_name = {"HK": "港股", "US": "美股"}.get(
                    top_market[0], top_market[0]
                )
                self.doc.add_paragraph(
                    f"分析：{market_name}贡献了最大利润（{float(top_market[1].net_profit):,.0f} HKD）。"
                )

        # 添加图表
        if "market_distribution" in charts and charts["market_distribution"]:
            self._add_image(charts["market_distribution"], "market_distribution")

    def _add_top_winners_section(self, stats: TradeStatistics) -> None:
        """添加最佳交易 Top 5 章节"""
        self.doc.add_heading("五、最佳交易 Top 5", level=1)

        if stats.top_winners:
            table = self.doc.add_table(rows=len(stats.top_winners) + 1, cols=6)
            table.style = "Table Grid"

            headers = ["排名", "标的", "盈利额", "盈利率", "持仓天数", "卖出日期"]
            data = []
            for w in stats.top_winners:
                data.append(
                    (
                        f"#{w.rank}",
                        f"{w.stock_name or w.code}",
                        f"+{float(w.profit_loss):,.2f}",
                        f"+{float(w.profit_loss_ratio):.1%}",
                        f"{w.holding_days}天",
                        w.sell_date or "",
                    )
                )

            self._fill_table(table, headers, data)
        else:
            self.doc.add_paragraph("暂无盈利交易记录。")

    def _add_top_losers_section(self, stats: TradeStatistics) -> None:
        """添加最大亏损 Top 5 章节"""
        self.doc.add_heading("六、最大亏损 Top 5", level=1)

        if stats.top_losers:
            table = self.doc.add_table(rows=len(stats.top_losers) + 1, cols=6)
            table.style = "Table Grid"

            headers = ["排名", "标的", "亏损额", "亏损率", "持仓天数", "卖出日期"]
            data = []
            for l in stats.top_losers:
                data.append(
                    (
                        f"#{l.rank}",
                        f"{l.stock_name or l.code}",
                        f"{float(l.profit_loss):,.2f}",
                        f"{float(l.profit_loss_ratio):.1%}",
                        f"{l.holding_days}天",
                        l.sell_date or "",
                    )
                )

            self._fill_table(table, headers, data)
        else:
            self.doc.add_paragraph("暂无亏损交易记录。")

    def _add_stock_stats_section(self, stats: TradeStatistics) -> None:
        """添加交易标的统计 Top 10 章节"""
        self.doc.add_heading("七、交易标的统计 Top 10", level=1)

        calculator = StatisticsCalculator()
        top_stocks = calculator.get_top_traded_stocks(stats, 10)

        if top_stocks:
            table = self.doc.add_table(rows=len(top_stocks) + 1, cols=5)
            table.style = "Table Grid"

            headers = ["标的", "交易次数", "胜率", "净盈亏", "状态"]
            data = []
            for s in top_stocks:
                status = "盈利" if s.net_profit > 0 else "亏损"
                data.append(
                    (
                        s.stock_name or s.code,
                        f"{s.trade_count}次",
                        f"{s.win_rate:.1%}",
                        f"{float(s.net_profit):,.2f}",
                        status,
                    )
                )

            self._fill_table(table, headers, data)
        else:
            self.doc.add_paragraph("暂无交易记录。")

    def _add_profit_loss_distribution_section(
        self, stats: TradeStatistics, charts: dict[str, bytes]
    ) -> None:
        """添加盈亏率分布章节"""
        self.doc.add_heading("八、盈亏率分布", level=1)

        if stats.profit_loss_buckets:
            # 只显示有数据的区间
            non_empty_buckets = [b for b in stats.profit_loss_buckets if b.count > 0]

            if non_empty_buckets:
                table = self.doc.add_table(rows=len(non_empty_buckets) + 1, cols=3)
                table.style = "Table Grid"

                headers = ["盈亏率区间", "交易笔数", "占比"]
                total = sum(b.count for b in non_empty_buckets)
                data = []
                for b in non_empty_buckets:
                    data.append(
                        (
                            b.bucket_name,
                            f"{b.count}笔",
                            f"{b.count/total:.1%}",
                        )
                    )

                self._fill_table(table, headers, data)

        # 添加图表
        if "profit_loss_bucket_bar" in charts and charts["profit_loss_bucket_bar"]:
            self._add_image(charts["profit_loss_bucket_bar"], "profit_loss_bucket_bar")

    def _add_monthly_section(
        self, stats: TradeStatistics, charts: dict[str, bytes]
    ) -> None:
        """添加月度盈亏趋势章节"""
        self.doc.add_heading("九、月度盈亏趋势", level=1)

        if stats.monthly_stats:
            table = self.doc.add_table(rows=len(stats.monthly_stats) + 1, cols=4)
            table.style = "Table Grid"

            headers = ["月份", "交易笔数", "胜率", "净盈亏"]
            data = []
            for month, ms in stats.monthly_stats.items():
                data.append(
                    (
                        month,
                        f"{ms.trade_count}笔",
                        f"{ms.win_rate:.1%}",
                        f"{float(ms.net_profit):,.2f}",
                    )
                )

            self._fill_table(table, headers, data)

        # 添加图表
        if "monthly_profit_bar" in charts and charts["monthly_profit_bar"]:
            self._add_image(charts["monthly_profit_bar"], "monthly_profit_bar")

    def _add_option_section(
        self, stats: TradeStatistics, option_trades: list[MatchedTrade]
    ) -> None:
        """添加期权交易统计章节"""
        self.doc.add_heading("十、期权交易统计", level=1)

        table = self.doc.add_table(rows=4, cols=3)
        table.style = "Table Grid"

        headers = ["指标", "数值", "说明"]
        data = [
            ("期权交易笔数", f"{stats.option_total_trades}笔", ""),
            ("期权胜率", f"{stats.option_win_rate:.1%}", ""),
            ("期权净盈亏", f"{float(stats.option_net_profit):,.2f}", ""),
        ]

        self._fill_table(table, headers, data)

        # 期权风险提示
        if stats.option_net_profit < 0:
            warning = (
                "风险提示：期权交易整体亏损，建议控制期权仓位比例（不超过总资产的15%），"
                "并严格执行止损策略。"
            )
            para = self.doc.add_paragraph(warning)
            para.runs[0].font.color.rgb = RGBColor(0x9C, 0x00, 0x06)

    def _add_conclusion_section(
        self,
        stats: TradeStatistics,
        stock_trades: list[MatchedTrade] = None,
        option_trades: list[MatchedTrade] = None,
    ) -> None:
        """
        添加结论与建议章节（基于投资框架 V10.10 的智能建议）

        使用 InvestmentCoach 生成专业的交易改进建议，
        参考框架: ~/Documents/trade/prompt/daily-analysis-prompt-v10_10.md
        """
        section_num = "十一" if stats.option_total_trades > 0 else "十"
        self.doc.add_heading(f"{section_num}、结论与建议", level=1)

        # 使用投资教练生成智能建议
        coach = InvestmentCoach()
        recommendation = coach.analyze(
            stats=stats,
            stock_trades=stock_trades or [],
            option_trades=option_trades or [],
        )

        # 添加框架版本说明
        framework_note = self.doc.add_paragraph(
            f"（基于投资分析框架 {recommendation.framework_version}）"
        )
        framework_note.runs[0].font.size = Pt(9)
        framework_note.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # === 风险警示（优先显示）===
        if recommendation.risk_alerts:
            self.doc.add_heading("⚠️ 风险警示", level=2)
            for alert in recommendation.risk_alerts:
                # 标题（加粗红色）
                para = self.doc.add_paragraph()
                title_run = para.add_run(f"• {alert.title}")
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)
                # 内容
                self.doc.add_paragraph(alert.content)
                # 框架参考
                if alert.framework_ref:
                    ref_para = self.doc.add_paragraph(f"📖 {alert.framework_ref}")
                    ref_para.runs[0].font.size = Pt(9)
                    ref_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # === 优势 ===
        self.doc.add_heading("✅ 优势", level=2)
        if recommendation.strengths:
            for strength in recommendation.strengths:
                para = self.doc.add_paragraph()
                title_run = para.add_run(f"• {strength.title}")
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                self.doc.add_paragraph(strength.content)
        else:
            self.doc.add_paragraph("暂无明显优势，建议持续优化交易策略。")

        # === 问题 ===
        self.doc.add_heading("❌ 需改进", level=2)
        if recommendation.weaknesses:
            for weakness in recommendation.weaknesses:
                para = self.doc.add_paragraph()
                title_run = para.add_run(f"• {weakness.title}")
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
                self.doc.add_paragraph(weakness.content)
                if weakness.framework_ref:
                    ref_para = self.doc.add_paragraph(f"📖 {weakness.framework_ref}")
                    ref_para.runs[0].font.size = Pt(9)
                    ref_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            self.doc.add_paragraph("整体表现良好，继续保持。")

        # === 建议 ===
        self.doc.add_heading("💡 改进建议", level=2)
        if recommendation.suggestions:
            for i, suggestion in enumerate(recommendation.suggestions, 1):
                # 标题（编号 + 加粗）
                para = self.doc.add_paragraph()
                title_run = para.add_run(f"{i}. {suggestion.title}")
                title_run.font.bold = True
                # 内容（处理多行）
                for line in suggestion.content.split("\n"):
                    if line.strip():
                        self.doc.add_paragraph(line)
                # 框架参考
                if suggestion.framework_ref:
                    ref_para = self.doc.add_paragraph(f"📖 {suggestion.framework_ref}")
                    ref_para.runs[0].font.size = Pt(9)
                    ref_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                self.doc.add_paragraph()  # 空行分隔
        else:
            self.doc.add_paragraph("继续保持当前策略。")

        # === 框架核心原则 ===
        self.doc.add_heading("📋 框架核心原则", level=2)
        principles = [
            "止损优先：股票 -10% 止损，期权 OCO 订单（+30%/-30%）",
            "估值先行：Forward PE + PB-ROE 双重筛选",
            "周期顺势：牛市满仓成长，熊市只做低估值",
            "量价确认：不追涨，等60分钟量价转换确认后入场",
            "完整计划：没有操作计划的交易 = 赌博",
        ]
        for p in principles:
            self.doc.add_paragraph(f"• {p}")

    def _fill_table(self, table, headers: list[str], data: list[tuple]) -> None:
        """填充表格数据"""
        # 表头
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # 表头样式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True

        # 数据行
        for row_idx, row_data in enumerate(data, start=1):
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                for col_idx, value in enumerate(row_data):
                    if col_idx < len(row.cells):
                        cell = row.cells[col_idx]
                        cell.text = str(value)
                        # 数据对齐
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_image(self, image_data: bytes, chart_key: str) -> None:
        """添加图片，使用双语图注"""
        if not image_data:
            return

        # 将字节数据转换为流
        image_stream = io.BytesIO(image_data)

        # 添加图片
        self.doc.add_picture(image_stream, width=Inches(6))

        # 图片居中
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加双语图注 (英文 + 中文)
        bilingual_caption = ChartGenerator.get_chart_caption(chart_key)
        caption_para = self.doc.add_paragraph(f"图：{bilingual_caption}")
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.runs[0].font.size = Pt(9)
        caption_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
